"""Word document generator using python-docx — invoices, contracts, reports."""

from __future__ import annotations

import io
import logging
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Brand colors
TESLE_GOLD = RGBColor(0xD4, 0xA8, 0x53)
TESLE_DARK = RGBColor(0x1A, 0x1A, 0x2E)
TESLE_TEXT = RGBColor(0x33, 0x33, 0x33)


class DOCXGenerator:
    """Generate Word documents with Tesle branding."""

    def generate(
        self,
        data: dict[str, Any],
        doc_type: str = "document",
        page_size: str = "a4",
        orientation: str = "portrait",
    ) -> bytes:
        """Generate a DOCX document."""
        doc = Document()

        if orientation == "landscape":
            for section in doc.sections:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width, section.page_height = (
                    section.page_height,
                    section.page_width,
                )

        if doc_type == "invoice":
            self._build_invoice(doc, data)
        elif doc_type == "contract":
            self._build_contract(doc, data)
        elif doc_type == "report":
            self._build_report(doc, data)
        else:
            self._build_generic(doc, data)

        buf = io.BytesIO()
        doc.save(buf)
        raw = buf.getvalue()
        buf.close()
        return raw

    # ── Invoice ──────────────────────────────────────────────────────────

    def _build_invoice(self, doc: Document, data: dict) -> None:
        company = data.get("company", {})

        # Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(company.get("name", "TESLE"))
        run.font.size = Pt(22)
        run.font.color.rgb = TESLE_GOLD
        run.bold = True

        if company.get("address"):
            doc.add_paragraph(company["address"])
        if company.get("email") or company.get("phone"):
            doc.add_paragraph(" | ".join(filter(None, [company.get("email"), company.get("phone")])))

        doc.add_paragraph()  # spacer

        # Invoice info
        p = doc.add_paragraph()
        run = p.add_run("Invoice #: ")
        run.bold = True
        p.add_run(data.get("invoice_number", "INV-0000"))

        p = doc.add_paragraph()
        run = p.add_run("Date: ")
        run.bold = True
        p.add_run(data.get("date", "N/A"))

        p = doc.add_paragraph()
        run = p.add_run("Due Date: ")
        run.bold = True
        p.add_run(data.get("due_date", "N/A"))

        p = doc.add_paragraph()
        run = p.add_run("Status: ")
        run.bold = True
        status = data.get("status", "pending")
        p.add_run(status.upper())

        doc.add_paragraph()

        # Bill To
        bill_to = data.get("bill_to", {})
        if bill_to:
            p = doc.add_paragraph()
            run = p.add_run("Bill To:")
            run.bold = True
            doc.add_paragraph(bill_to.get("name", ""))
            if bill_to.get("address"):
                doc.add_paragraph(bill_to["address"])
            if bill_to.get("email"):
                doc.add_paragraph(bill_to["email"])

        doc.add_paragraph()

        # Line items table
        items = data.get("items", [])
        if items:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            headers = ["#", "Description", "Qty", "Unit Price", "Amount"]
            for i, h in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = h
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)

            for idx, item in enumerate(items, 1):
                qty = item.get("quantity", 1)
                price = float(item.get("unit_price", item.get("price", 0)))
                row = table.add_row()
                row.cells[0].text = str(idx)
                row.cells[1].text = item.get("description", item.get("name", ""))
                row.cells[2].text = str(qty)
                row.cells[3].text = f"${price:,.2f}"
                row.cells[4].text = f"${qty * price:,.2f}"

        doc.add_paragraph()

        # Totals
        subtotal = data.get("subtotal", 0)
        tax_rate = data.get("tax_rate", 0)
        total = data.get("total", 0)

        if not total and items:
            subtotal = sum(
                float(it.get("quantity", 1)) * float(it.get("unit_price", it.get("price", 0)))
                for it in items
            )
            total = subtotal + subtotal * (tax_rate / 100) if tax_rate else subtotal

        if subtotal:
            p = doc.add_paragraph()
            run = p.add_run("Subtotal: ")
            run.bold = True
            p.add_run(f"${subtotal:,.2f}")

        if tax_rate:
            p = doc.add_paragraph()
            run = p.add_run(f"Tax ({tax_rate}%): ")
            run.bold = True
            p.add_run(f"${subtotal * tax_rate / 100:,.2f}")

        p = doc.add_paragraph()
        run = p.add_run("Total: ")
        run.bold = True
        run.font.size = Pt(12)
        p.add_run(f"${total:,.2f}")

        # Notes
        if data.get("notes"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run("Notes:")
            run.bold = True
            doc.add_paragraph(data["notes"])

    # ── Contract ─────────────────────────────────────────────────────────

    def _build_contract(self, doc: Document, data: dict) -> None:
        # Title
        title = data.get("title", "Contract Agreement")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = TESLE_DARK

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Effective Date: {data.get('effective_date', 'N/A')}")

        doc.add_paragraph()

        # Parties
        heading = doc.add_heading("Parties", level=2)
        for run in heading.runs:
            run.font.color.rgb = TESLE_DARK

        parties = data.get("parties", [])
        for party in parties:
            p = doc.add_paragraph()
            run = p.add_run(f"{party.get('role', 'Party')}: ")
            run.bold = True
            p.add_run(party.get("name", ""))
            if party.get("address"):
                doc.add_paragraph(party["address"])

        doc.add_paragraph()

        # Clauses
        clauses = data.get("clauses", [])
        for i, clause in enumerate(clauses, 1):
            heading = doc.add_heading(f"{i}. {clause.get('heading', f'Section {i}')}", level=2)
            for run in heading.runs:
                run.font.color.rgb = TESLE_DARK
            doc.add_paragraph(clause.get("content", ""))

        # Rendered content
        if data.get("_rendered_content"):
            doc.add_paragraph()
            for line in data["_rendered_content"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

        # Signatures
        doc.add_paragraph()
        doc.add_paragraph()
        sig_parties = parties if len(parties) >= 2 else [
            {"name": data.get("party_a_name", "Party A"), "role": "Authorized Signatory"},
            {"name": data.get("party_b_name", "Party B"), "role": "Authorized Signatory"},
        ]
        for party in sig_parties:
            doc.add_paragraph("_" * 40)
            p = doc.add_paragraph()
            run = p.add_run(party.get("name", ""))
            run.bold = True
            doc.add_paragraph(party.get("role", ""))
            doc.add_paragraph("Date: ___________________")

    # ── Report ───────────────────────────────────────────────────────────

    def _build_report(self, doc: Document, data: dict) -> None:
        # Title page
        for _ in range(6):
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get("title", "Business Report"))
        run.font.size = Pt(26)
        run.bold = True
        run.font.color.rgb = TESLE_GOLD

        if data.get("subtitle"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(data["subtitle"])

        if data.get("date"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Date: {data['date']}")

        if data.get("author"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(f"Prepared by: {data['author']}")

        doc.add_page_break()

        # Summary
        if data.get("summary"):
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(data["summary"])

        # Sections
        for section in data.get("sections", []):
            doc.add_heading(section.get("heading", ""), level=1)
            content = section.get("content", "")
            for para in content.split("\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        # Data tables
        for tbl in data.get("tables", []):
            if tbl.get("title"):
                doc.add_heading(tbl["title"], level=2)
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Light Grid Accent 1"
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                for row_data in rows:
                    row = table.add_row()
                    for i, val in enumerate(row_data):
                        row.cells[i].text = str(val)

        # Rendered content
        if data.get("_rendered_content"):
            doc.add_paragraph()
            for line in data["_rendered_content"].split("\n"):
                if line.strip():
                    doc.add_paragraph(line)

    # ── Generic ──────────────────────────────────────────────────────────

    def _build_generic(self, doc: Document, data: dict) -> None:
        if data.get("title"):
            doc.add_heading(data["title"], level=0)
        if data.get("subtitle"):
            doc.add_heading(data["subtitle"], level=1)
        content = data.get("content", data.get("_rendered_content", ""))
        if content:
            for line in content.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
